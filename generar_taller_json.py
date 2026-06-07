from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


BASE_DIR = Path(__file__).resolve().parent
OUTPUT = BASE_DIR / "Taller_Consumo_JSON_LoginApp.docx"
SCREENSHOT = BASE_DIR / "captura_usuarios_json_apa.png"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def apply_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_paragraph(paragraph, size=11, bold=False, color="000000", align=None):
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(6)
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(14 if level == 1 else 12)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 2.0
    return paragraph


def add_body(doc, text, bold=False):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    style_paragraph(paragraph, bold=bold)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    style_paragraph(paragraph)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    style_paragraph(paragraph)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    return paragraph


def add_reference(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    style_paragraph(paragraph)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    return paragraph


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.line_spacing = 2.0
title.paragraph_format.space_after = Pt(24)
run = title.add_run("Taller: Consumo JSON en la aplicación Android LoginApp")
run.font.name = "Calibri"
run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
run.font.size = Pt(16)
run.font.bold = True

for line in [
    "Actividad: Consumo JSON de un webservice externo",
    "Aplicación: LoginApp",
    "Estudiante: ______________________________",
    "Asignatura: ______________________________",
    "Docente: ______________________________",
    "Fecha: 7 de junio de 2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.add_run(line)
    style_paragraph(p)

doc.add_page_break()

add_heading(doc, "Resumen", 1)
add_body(
    doc,
    "En este taller se documenta la implementación del consumo de datos JSON en la aplicación Android LoginApp. "
    "La actividad consistió en agregar una pantalla nueva que consulta un webservice externo, procesa la respuesta JSON "
    "y muestra información de usuarios dentro de la interfaz de la aplicación. La solución fue desarrollada en Java, "
    "utilizando HttpURLConnection para realizar la petición HTTP y las clases JSONArray y JSONObject para leer los datos recibidos."
)

add_heading(doc, "Introducción", 1)
add_body(
    doc,
    "El consumo de servicios web es una función común en aplicaciones móviles modernas, ya que permite que una app obtenga "
    "información desde fuentes externas y la presente al usuario en tiempo real. En esta actividad se integró el servicio "
    "JSONPlaceholder, el cual entrega datos de prueba en formato JSON. La implementación permite evidenciar el uso de permisos "
    "de internet, la ejecución de tareas de red fuera del hilo principal y la presentación ordenada de la información recibida."
)

add_heading(doc, "Objetivos", 1)
add_body(doc, "Objetivo general", bold=True)
add_body(
    doc,
    "Implementar el consumo de un archivo JSON proveniente de un webservice externo en la aplicación Android LoginApp."
)
add_body(doc, "Objetivos específicos", bold=True)
for item in [
    "Agregar el permiso de internet necesario para realizar peticiones HTTP desde Android.",
    "Crear una pantalla que consuma datos JSON del endpoint de usuarios de JSONPlaceholder.",
    "Procesar la respuesta JSON para extraer nombre, correo, ciudad y empresa de cada usuario.",
    "Mostrar la información obtenida en una interfaz clara y verificable dentro de la aplicación.",
]:
    add_bullet(doc, item)

add_heading(doc, "Desarrollo de la actividad", 1)
add_heading(doc, "Descripción de la aplicación", 2)
add_body(
    doc,
    "LoginApp era una aplicación base con pantallas de inicio de sesión y registro. Para completar la actividad se agregó "
    "una opción llamada Ver usuarios JSON en la pantalla principal. Esta opción abre una nueva pantalla llamada UsersActivity, "
    "donde se realiza el consumo del servicio externo y se muestra el resultado."
)

add_heading(doc, "Cambios realizados", 2)
for item in [
    "Se agregó el permiso <uses-permission android:name=\"android.permission.INTERNET\" /> en AndroidManifest.xml.",
    "Se creó UsersActivity.java para manejar la petición al webservice, la lectura del JSON y la actualización de la interfaz.",
    "Se creó activity_users.xml con título, botón Actualizar, mensaje de estado y área para listar usuarios.",
    "Se agregó el botón Ver usuarios JSON en activity_main.xml y su evento de navegación en MainActivity.java.",
]:
    add_bullet(doc, item)

add_heading(doc, "Flujo de funcionamiento", 2)
for item in [
    "El usuario abre LoginApp desde el dispositivo Android.",
    "Desde la pantalla principal selecciona la opción Ver usuarios JSON.",
    "La aplicación abre UsersActivity y ejecuta la consulta a https://jsonplaceholder.typicode.com/users.",
    "El servicio devuelve un arreglo JSON con 10 usuarios.",
    "La aplicación lee cada objeto JSON y presenta nombre, correo, ciudad y empresa.",
    "El botón Actualizar permite ejecutar nuevamente el consumo del webservice.",
]:
    add_numbered(doc, item)

add_heading(doc, "Resumen técnico", 2)
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
apply_table_borders(table)
hdr = table.rows[0].cells
hdr[0].text = "Elemento"
hdr[1].text = "Descripción"
for cell in hdr:
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        style_paragraph(p, bold=True)

rows = [
    ("Webservice", "JSONPlaceholder: endpoint /users."),
    ("Formato de datos", "JSON, recibido como un arreglo de objetos."),
    ("Clase principal", "UsersActivity.java."),
    ("Conexión HTTP", "HttpURLConnection con método GET."),
    ("Procesamiento JSON", "JSONArray y JSONObject."),
    ("Datos mostrados", "Nombre, correo, ciudad y empresa."),
]
for left, right in rows:
    cells = table.add_row().cells
    cells[0].text = left
    cells[1].text = right
    for cell in cells:
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            style_paragraph(p)
set_table_widths(table, [1.8, 4.7])

add_heading(doc, "Evidencia de funcionamiento", 1)
add_body(
    doc,
    "La siguiente figura evidencia que la pantalla Usuarios desde JSON cargó correctamente los datos del webservice externo. "
    "Se observa el mensaje de confirmación y la lista de usuarios recibidos."
)
if SCREENSHOT.exists():
    pic_paragraph = doc.add_paragraph()
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_paragraph.add_run()
    run.add_picture(str(SCREENSHOT), width=Inches(3.0))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 2.0
    caption.add_run("Figura 1\nPantalla Usuarios desde JSON con datos cargados desde JSONPlaceholder.")
    style_paragraph(caption)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("Nota. Captura tomada desde el dispositivo Android conectado durante la prueba de la aplicación.")
    style_paragraph(note, size=10)
else:
    add_body(doc, "No se encontró la captura de pantalla para anexarla al documento.")

add_heading(doc, "Resultados", 1)
for item in [
    "La aplicación se instaló correctamente en el dispositivo Android conectado.",
    "La pantalla Usuarios desde JSON abrió desde el botón agregado en la pantalla principal.",
    "El servicio externo respondió correctamente y entregó 10 registros de usuarios.",
    "La app presentó la información en pantalla de forma clara y ordenada.",
]:
    add_bullet(doc, item)

add_heading(doc, "Conclusiones", 1)
add_body(
    doc,
    "La actividad permitió completar la integración de un consumo JSON real dentro de la aplicación LoginApp. "
    "El desarrollo demuestra que la aplicación puede comunicarse con un servicio externo, interpretar una respuesta JSON "
    "y transformar los datos en información visible para el usuario. Además, se evidencia la importancia de declarar el permiso "
    "de internet en el manifiesto de Android y de ejecutar las operaciones de red en un hilo secundario para evitar bloqueos en la interfaz."
)

add_heading(doc, "Referencias", 1)
references = [
    "Android Developers. (s. f.). HttpURLConnection. https://developer.android.com/reference/java/net/HttpURLConnection",
    "Android Developers. (s. f.). Manifest.permission. https://developer.android.com/reference/android/Manifest.permission",
    "Android Developers. (s. f.). App manifest overview. https://developer.android.com/guide/topics/manifest/manifest-intro",
    "JSONPlaceholder. (s. f.). JSONPlaceholder: Free fake and reliable API for testing and prototyping. https://jsonplaceholder.typicode.com/",
]
for ref in references:
    add_reference(doc, ref)

doc.save(OUTPUT)
print(OUTPUT)
